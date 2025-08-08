# # processing.py (Corrected)
# import requests
# from pypdf import PdfReader
# from io import BytesIO


# # Import the correct and robust text splitter
# from langchain_text_splitters import RecursiveCharacterTextSplitter


# def load_document_from_url(url: str) -> str:
#     """Downloads a PDF from a URL and extracts its text."""
#     try:
#         response = requests.get(url)
#         response.raise_for_status()  # Raise an exception for bad status codes

#         # Read PDF from in-memory content
#         pdf_file = BytesIO(response.content)
#         reader = PdfReader(pdf_file)
#         text = ""
#         for page in reader.pages:
#             # Add a space between pages to ensure sentences don't merge
#             text += (page.extract_text() or "") + " "
#         return text
#     except requests.RequestException as e:
#         print(f"Error downloading file: {e}")
#         return None
#     except Exception as e:
#         print(f"Error processing PDF: {e}")
#         return None


# def chunk_text(text: str) -> list[str]:
#     """Splits text into semantic chunks using a robust method."""
    
#     # This is the industry-standard way to create chunks for RAG
#     text_splitter = RecursiveCharacterTextSplitter(
#         chunk_size=1000,  # The maximum size of each chunk (in characters)
#         chunk_overlap=100,  # The number of characters to overlap between chunks
#         length_function=len,
#     )
    
#     # The split_text method does the work
#     return text_splitter.split_text(text)


# processing.py (Corrected and Professional)
import requests
from pypdf import PdfReader
from docx import Document
from io import BytesIO
import re
import logging
from langchain_text_splitters import RecursiveCharacterTextSplitter
import email
from email.policy import default

# --- ADD THESE IMPORTS AT THE TOP ---
from urllib.parse import urlparse
import os

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def _load_pdf_text(content: bytes) -> str:
    """Extracts clean, properly formatted text from PDF content."""
    try:
        pdf_file = BytesIO(content)
        reader = PdfReader(pdf_file)
        
        full_text = ""
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                cleaned_text = _clean_text(page_text)
                if cleaned_text:
                    full_text += f"[Page {page_num + 1}] {cleaned_text}\n\n"
        
        logger.info(f"Successfully extracted text from {len(reader.pages)} PDF pages")
        return full_text.strip()
        
    except Exception as e:
        logger.error(f"Error processing PDF: {e}")
        return None


def _load_docx_text(content: bytes) -> str:
    """Extracts text from a DOCX document."""
    try:
        docx_file = BytesIO(content)
        doc = Document(docx_file)
        
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        logger.info("Successfully extracted text from DOCX document")
        return "\n\n".join(full_text)
        
    except Exception as e:
        logger.error(f"Error processing DOCX: {e}")
        return None


def _load_eml_text(content: bytes) -> str:
    """Extracts the body text from an EML email file."""
    try:
        msg = email.message_from_bytes(content, policy=default)
        
        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                content_disposition = str(part.get("Content-Disposition"))

                if content_type == "text/plain" and "attachment" not in content_disposition:
                    body = part.get_payload(decode=True).decode()
                    break
        else:
            body = msg.get_payload(decode=True).decode()
            
        logger.info("Successfully extracted text from EML file")
        return body

    except Exception as e:
        logger.error(f"Error processing EML file: {e}")
        return None


def load_document_from_url(url: str) -> str:
    """
    Downloads a document (PDF, DOCX, or EML) from a URL and extracts clean text.
    """
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        content = response.content

        parsed_url = urlparse(url)
        path = parsed_url.path
        
        _, file_extension_with_dot = os.path.splitext(path)
        file_extension = file_extension_with_dot.lower().strip('.') 


        if file_extension == 'pdf':
            return _load_pdf_text(content)
        elif file_extension == 'docx':
            return _load_docx_text(content)
        elif file_extension == 'eml':
            return _load_eml_text(content)
        else:
            # This is the error message you were seeing
            logger.error(f"Unsupported file type: {file_extension}")
            return None

    except requests.RequestException as e:
        logger.error(f"Error downloading file: {e}")
        return None
    except Exception as e:
        logger.error(f"An unexpected error occurred: {e}")
        return None


def _clean_text(text: str) -> str:
    """Clean and normalize extracted text for better processing."""
    if not text:
        return ""
    
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    text = re.sub(r'(\.)([A-Z])', r'\1 \2', text)
    text = re.sub(r'Page \d+.*?\n', '', text)
    text = re.sub(r'\n\d+\n', '\n', text)
    text = re.sub(r'([.!?])\s*([.!?])+', r'\1', text)
    text = re.sub(r'\s+([,.!?;:])', r'\1', text)
    text = re.sub(r'([.!?;:])\s*([A-Z])', r'\1 \2', text)
    
    return text.strip()


def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> list[str]:
    """Splits text into semantic chunks."""
    logger.info(f"Chunking text with size={chunk_size}, overlap={chunk_overlap}")
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " ", ""],
        keep_separator=True
    )
    
    chunks = text_splitter.split_text(text)
    cleaned_chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
    
    logger.info(f"Successfully created {len(cleaned_chunks)} text chunks")
    return cleaned_chunks


def get_text_stats(text: str) -> dict:
    """Get statistics about the extracted text."""
    if not text:
        return {"error": "No text provided"}
    
    words = text.split()
    sentences = re.split(r'[.!?]+', text)
    paragraphs = text.split('\n\n')
    
    stats = {
        "total_characters": len(text),
        "total_words": len(words),
        "total_sentences": len([s for s in sentences if s.strip()]),
        "total_paragraphs": len([p for p in paragraphs if p.strip()]),
        "average_words_per_sentence": len(words) / max(len([s for s in sentences if s.strip()]), 1),
        "has_policy_terms": bool(re.search(r'\b(policy|premium|coverage|claim|benefit|deductible)\b', text, re.IGNORECASE)),
        "has_dates": bool(re.search(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', text)),
        "has_currency": bool(re.search(r'[\$₹£€]\s*\d+|Rs\.?\s*\d+', text, re.IGNORECASE))
    }
    
    return stats